import docx
# download example document from
# http://autbor.com/demo.docx

d = docx.Document('c:\\users\\mkhir\\documents\\demo.docx')

# viewing:
print(d.paragraphs)
print(d.paragraphs[0].text) # Document Title
p = print(d.paragraphs[1].text) # 'A plain paragraph having some bold and some italic.'

print(p.runs)
print(p.runs[0].text) # 'A plain paragraph having some '
print(p.runs[1].text) # 'bold'
print(p.runs[2].text) # ' and some '
print(p.runs[3].text) # 'italic.'

print(p.runs[1].bold) # True
print(p.runs[3].italic) # True

# editing and saving:
p.runs[3].underline = True
p.runs[3].text = 'italic and underlined.'
d.save('c:\\demo2.docx')

p.style = 'Title'
d.save('c:\\demo3.docx')

# creating and saving new blank document:
d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('This is another paragraph.')
d.save('c:\\demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True
d.save('c:\\demo5.docx')

# print all text in a document
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)

print(getText('c:\\users\\mkhir\\documents\\demo.docx'))